# import docx
# from data_extractor.file_loaders.file_loader import FileLoader

# class DOCXLoader(FileLoader):
#     def validate_file(self, file_path: str) -> bool:
#         return file_path.lower().endswith('.docx')

#     def load_file(self, file_path: str) -> docx.Document:
#         if not self.validate_file(file_path):
#             raise ValueError("Invalid DOCX file.")
#         try:
#             # Attempt to open the DOCX file
#             return docx.Document(file_path)
#         except Exception:
#             # Catch any exception that occurs and raise a ValueError
#             raise ValueError("Invalid DOCX file.")




import docx
import msoffcrypto
from data_extractor.file_loaders.file_loader import FileLoader
from io import BytesIO
from docx.opc.exceptions import PackageNotFoundError

class DOCXLoader(FileLoader):
    def validate_file(self, file_path: str) -> bool:
        return file_path.lower().endswith('.docx')

    def load_file(self, file_path: str) -> docx.Document:
        if not self.validate_file(file_path):
            raise ValueError("Invalid DOCX file.")

        try:
            # Attempt to open the DOCX file directly
            return docx.Document(file_path)

        except PackageNotFoundError:
            # This exception means the file is likely password-protected
            with open(file_path, 'rb') as f:
                office_file = msoffcrypto.OfficeFile(f)

                if office_file.is_encrypted:
                    # Prompt for password and decrypt
                    password = input("Enter your DOCX password: ")
                    try:
                        office_file.load_key(password=password)
                    except msoffcrypto.exceptions.InvalidKeyError:
                        raise ValueError("Failed to decrypt the DOCX file: Incorrect password.")

                    decrypted_file = BytesIO()
                    office_file.decrypt(decrypted_file)  # Decrypt the file
                    decrypted_file.seek(0)  # Reset the BytesIO cursor to the beginning

                    try:
                        return docx.Document(decrypted_file)  # Load the decrypted document
                    except Exception as e:
                        raise ValueError(f"Failed to load decrypted DOCX file: {str(e)}")
                else:
                    raise ValueError("The DOCX file is reported as unencrypted, but a password is required.")

        except Exception as e:
            # Catch any other exceptions and raise a ValueError with details
            raise ValueError(f"Invalid DOCX file: {str(e)}")
